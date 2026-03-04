# -*- coding: utf-8 -*-
"""
Word 문서 생성기 (generate_doc.py)

JSON 템플릿을 기반으로 서식이 적용된 .docx 파일을 생성한다.

사용법:
    # Mode A: 원본 재현
    python generate_doc.py <template.json> <output.docx>

    # Mode B: 새 콘텐츠 적용
    python generate_doc.py <template.json> <content.json> <output.docx>

필수 환경:
    - Python 3.8+
    - python-docx (pip install python-docx)
"""

import json
import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.section import WD_ORIENT, WD_SECTION_START
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("오류: python-docx가 설치되지 않았습니다. pip install python-docx")
    sys.exit(1)

VERSION = "1.0.0"


# ─── 유틸리티 함수 ──────────────────────────────────────

def hex_to_rgb_color(hex_str):
    """#RRGGBB → RGBColor 객체"""
    if not hex_str:
        return None
    hex_str = hex_str.lstrip("#")
    if len(hex_str) != 6:
        return None
    r = int(hex_str[0:2], 16)
    g = int(hex_str[2:4], 16)
    b = int(hex_str[4:6], 16)
    return RGBColor(r, g, b)


def cm_val(cm_float):
    """cm float → python-docx Cm 단위"""
    if cm_float is None or cm_float == 0:
        return None
    return Cm(cm_float)


ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "distribute": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
}

LINE_SPACING_MAP = {
    "single": WD_LINE_SPACING.SINGLE,
    "1.5lines": WD_LINE_SPACING.ONE_POINT_FIVE,
    "double": WD_LINE_SPACING.DOUBLE,
    "exactly": WD_LINE_SPACING.EXACTLY,
    "atLeast": WD_LINE_SPACING.AT_LEAST,
    "multiple": WD_LINE_SPACING.MULTIPLE,
}

VERTICAL_ALIGNMENT_MAP = {
    "top": WD_CELL_VERTICAL_ALIGNMENT.TOP,
    "center": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
    "bottom": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
}


def str_to_alignment(alignment_str):
    return ALIGNMENT_MAP.get(alignment_str, WD_ALIGN_PARAGRAPH.LEFT)


# ─── DocGenerator 클래스 ────────────────────────────────

class DocGenerator:
    """JSON 템플릿 기반 Word 문서 생성기"""

    def __init__(self):
        self.template = None
        self.content_data = None
        self.doc = None
        self.base_dir = "."

    def load_template(self, template_path):
        """JSON 템플릿 로드 및 검증"""
        with open(template_path, "r", encoding="utf-8") as f:
            self.template = json.load(f)
        self.base_dir = os.path.dirname(os.path.abspath(template_path))
        self._validate_template()

    def load_content(self, content_path):
        """새 콘텐츠 데이터 로드 (Mode B)"""
        with open(content_path, "r", encoding="utf-8") as f:
            self.content_data = json.load(f)

    def generate(self):
        """문서 생성 실행"""
        self.doc = Document()

        # 페이지 설정
        self._setup_page(self.template.get("page_setup", {}))

        # 다중 섹션 설정
        sections = self.template.get("sections", [])
        for i, section_def in enumerate(sections):
            if i > 0:
                new_section = self.doc.add_section(WD_SECTION_START.NEW_PAGE)
                self._setup_section(new_section, section_def)
            else:
                self._setup_section(self.doc.sections[0], section_def)

        # 스타일 등록
        self._setup_styles()

        # 콘텐츠 생성
        self._generate_content()

        # 머리글/바닥글 설정
        self._setup_headers_footers()

    def save(self, output_path):
        """문서 저장"""
        os.makedirs(os.path.dirname(os.path.abspath(output_path)) or ".", exist_ok=True)
        self.doc.save(output_path)
        return output_path

    # ─── 내부 메서드 ─────────────────────────────────────

    def _validate_template(self):
        """템플릿 필수 필드 검증"""
        required = ["meta", "page_setup", "styles"]
        for field in required:
            if field not in self.template:
                raise ValueError(f"필수 필드 누락: {field}")

    def _setup_page(self, page_setup):
        """기본 페이지 설정"""
        section = self.doc.sections[0]

        if page_setup.get("page_width_cm"):
            section.page_width = Cm(page_setup["page_width_cm"])
        if page_setup.get("page_height_cm"):
            section.page_height = Cm(page_setup["page_height_cm"])

        if page_setup.get("orientation") == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            # 가로 방향 시 width/height 교환
            if section.page_width < section.page_height:
                section.page_width, section.page_height = (
                    section.page_height,
                    section.page_width,
                )

        margin_keys = {
            "margin_top_cm": "top_margin",
            "margin_bottom_cm": "bottom_margin",
            "margin_left_cm": "left_margin",
            "margin_right_cm": "right_margin",
        }
        for json_key, attr in margin_keys.items():
            if page_setup.get(json_key) is not None:
                setattr(section, attr, Cm(page_setup[json_key]))

        if page_setup.get("header_distance_cm") is not None:
            section.header_distance = Cm(page_setup["header_distance_cm"])
        if page_setup.get("footer_distance_cm") is not None:
            section.footer_distance = Cm(page_setup["footer_distance_cm"])
        if page_setup.get("gutter_cm") is not None:
            section.gutter = Cm(page_setup["gutter_cm"])

    def _setup_section(self, section, section_def):
        """개별 섹션 설정"""
        ps = section_def.get("page_setup")
        if ps:
            self._setup_page_for_section(section, ps)

        if section_def.get("has_different_first_page"):
            section.different_first_page_header_footer = True

    def _setup_page_for_section(self, section, page_setup):
        """섹션별 페이지 설정 적용"""
        if page_setup.get("page_width_cm"):
            section.page_width = Cm(page_setup["page_width_cm"])
        if page_setup.get("page_height_cm"):
            section.page_height = Cm(page_setup["page_height_cm"])
        if page_setup.get("orientation") == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            if section.page_width < section.page_height:
                section.page_width, section.page_height = (
                    section.page_height,
                    section.page_width,
                )

        for json_key, attr in {
            "margin_top_cm": "top_margin",
            "margin_bottom_cm": "bottom_margin",
            "margin_left_cm": "left_margin",
            "margin_right_cm": "right_margin",
        }.items():
            if page_setup.get(json_key) is not None:
                setattr(section, attr, Cm(page_setup[json_key]))

    def _setup_styles(self):
        """스타일 등록"""
        styles = self.template.get("styles", {})

        for style_def in styles.get("paragraph_styles", []):
            self._register_style(style_def, WD_STYLE_TYPE.PARAGRAPH)

        for style_def in styles.get("character_styles", []):
            self._register_style(style_def, WD_STYLE_TYPE.CHARACTER)

    def _register_style(self, style_def, style_type):
        """개별 스타일 등록"""
        name = style_def.get("name")
        if not name:
            return

        try:
            style = self.doc.styles.add_style(name, style_type)
        except ValueError:
            # 이미 존재하는 스타일
            try:
                style = self.doc.styles[name]
            except KeyError:
                return

        # 폰트 설정
        font_def = style_def.get("font", {})
        if font_def:
            self._apply_font(style.font, font_def)
            # 동아시아 글꼴 (XML)
            if font_def.get("name_east_asia"):
                self._set_east_asia_font(style, font_def["name_east_asia"])

        # 단락 서식 (단락 스타일만)
        if style_type == WD_STYLE_TYPE.PARAGRAPH:
            pf_def = style_def.get("paragraph_format", {})
            if pf_def:
                self._apply_paragraph_format(style.paragraph_format, pf_def)

    def _apply_font(self, font, font_def):
        """폰트 속성 적용"""
        if font_def.get("name_ascii"):
            font.name = font_def["name_ascii"]
        if font_def.get("size_pt"):
            font.size = Pt(font_def["size_pt"])
        if font_def.get("bold") is not None:
            font.bold = font_def["bold"]
        if font_def.get("italic") is not None:
            font.italic = font_def["italic"]
        if font_def.get("color"):
            rgb = hex_to_rgb_color(font_def["color"])
            if rgb:
                font.color.rgb = rgb
        if font_def.get("underline") and font_def["underline"] != "none":
            font.underline = True
        if font_def.get("strikethrough"):
            font.strike = font_def["strikethrough"]
        if font_def.get("superscript") is not None:
            font.superscript = font_def["superscript"]
        if font_def.get("subscript") is not None:
            font.subscript = font_def["subscript"]
        if font_def.get("all_caps") is not None:
            font.all_caps = font_def["all_caps"]

    def _set_east_asia_font(self, style_or_run, font_name):
        """동아시아 글꼴 설정 (w:eastAsia XML 직접 조작)"""
        element = (
            style_or_run.element
            if hasattr(style_or_run, "element")
            else style_or_run._element
        )
        rPr = element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), font_name)

    def _apply_paragraph_format(self, pf, pf_def):
        """단락 서식 적용"""
        if pf_def.get("alignment"):
            pf.alignment = str_to_alignment(pf_def["alignment"])

        if pf_def.get("line_spacing_rule"):
            rule_str = pf_def["line_spacing_rule"]
            rule = LINE_SPACING_MAP.get(rule_str, WD_LINE_SPACING.SINGLE)
            pf.line_spacing_rule = rule
            if pf_def.get("line_spacing"):
                val = pf_def["line_spacing"]
                if rule == WD_LINE_SPACING.MULTIPLE:
                    pf.line_spacing = val
                elif rule in (WD_LINE_SPACING.EXACTLY, WD_LINE_SPACING.AT_LEAST):
                    pf.line_spacing = Pt(val)

        if pf_def.get("space_before_pt") is not None:
            pf.space_before = Pt(pf_def["space_before_pt"])
        if pf_def.get("space_after_pt") is not None:
            pf.space_after = Pt(pf_def["space_after_pt"])
        if pf_def.get("first_line_indent_cm") is not None:
            pf.first_line_indent = Cm(pf_def["first_line_indent_cm"])
        if pf_def.get("left_indent_cm") is not None:
            pf.left_indent = Cm(pf_def["left_indent_cm"])
        if pf_def.get("right_indent_cm") is not None:
            pf.right_indent = Cm(pf_def["right_indent_cm"])
        if pf_def.get("keep_with_next") is not None:
            pf.keep_with_next = pf_def["keep_with_next"]
        if pf_def.get("keep_together") is not None:
            pf.keep_together = pf_def["keep_together"]
        if pf_def.get("page_break_before") is not None:
            pf.page_break_before = pf_def["page_break_before"]

    def _generate_content(self):
        """콘텐츠 생성"""
        # Mode B: 새 콘텐츠 데이터 사용
        if self.content_data and "content" in self.content_data:
            content_items = self.content_data["content"]
        # Mode A: 템플릿 콘텐츠 재현
        elif "content" in self.template:
            content_items = self.template["content"]
        else:
            return

        for item in content_items:
            item_type = item.get("type")
            if item_type == "paragraph":
                self._add_paragraph(item)
            elif item_type == "table":
                self._add_table(item)
            elif item_type == "image":
                self._add_image(item)
            elif item_type == "page_break":
                self.doc.add_page_break()

    def _add_paragraph(self, para_def):
        """서식 적용된 단락 추가"""
        style_name = para_def.get("style")

        # 스타일 존재 확인
        try:
            style = self.doc.styles[style_name] if style_name else None
        except KeyError:
            style = None

        p = self.doc.add_paragraph(style=style)

        # runs가 있으면 run별로 추가
        runs = para_def.get("runs", [])
        if runs:
            for run_def in runs:
                run = p.add_run(run_def.get("text", ""))
                font_def = run_def.get("font", {})
                if font_def:
                    self._apply_font(run.font, font_def)
                    if font_def.get("name_east_asia"):
                        self._set_east_asia_font_run(run, font_def["name_east_asia"])
        else:
            # runs가 없으면 text를 직접 추가
            text = para_def.get("text", "")
            if text:
                p.add_run(text)

        # 단락 서식 오버라이드
        pf_def = para_def.get("paragraph_format")
        if pf_def:
            self._apply_paragraph_format(p.paragraph_format, pf_def)

        # 번호 매기기
        numbering = para_def.get("numbering")
        if numbering:
            self._apply_numbering(p, numbering)

    def _set_east_asia_font_run(self, run, font_name):
        """Run 레벨 동아시아 글꼴 설정"""
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), font_name)

    def _apply_numbering(self, paragraph, numbering_def):
        """번호 매기기 적용"""
        pPr = paragraph._p.get_or_add_pPr()
        numPr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), str(numbering_def.get("level", 0)))
        numId = OxmlElement("w:numId")
        numId.set(qn("w:val"), str(numbering_def.get("definition_id", 1)))
        numPr.append(ilvl)
        numPr.append(numId)
        pPr.append(numPr)

    def _add_table(self, table_def):
        """서식 적용된 표 추가"""
        rows = table_def.get("rows", 1)
        cols = table_def.get("cols", 1)

        table = self.doc.add_table(rows=rows, cols=cols)
        table.autofit = False

        # 열 너비 설정
        col_widths = table_def.get("col_widths_cm", [])
        for i, width in enumerate(col_widths):
            if i < len(table.columns):
                table.columns[i].width = Cm(width)

        # 셀 데이터 채우기
        data = table_def.get("data", [])
        for row_idx, row_data in enumerate(data):
            if row_idx >= rows:
                break
            for col_idx, cell_def in enumerate(row_data):
                if col_idx >= cols:
                    break
                cell = table.cell(row_idx, col_idx)

                # 텍스트
                text = cell_def.get("text", "") if isinstance(cell_def, dict) else str(cell_def)
                cell.text = text

                if isinstance(cell_def, dict):
                    # 셀 음영
                    shading_color = cell_def.get("shading_color")
                    if shading_color:
                        self._apply_cell_shading(cell, shading_color)

                    # 수직 정렬
                    v_align = cell_def.get("vertical_alignment")
                    if v_align:
                        cell.vertical_alignment = VERTICAL_ALIGNMENT_MAP.get(
                            v_align, WD_CELL_VERTICAL_ALIGNMENT.TOP
                        )

                    # 셀 폰트
                    font_def = cell_def.get("font", {})
                    if font_def and cell.paragraphs:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                self._apply_font(run.font, font_def)

                    # 셀 병합
                    merge = cell_def.get("merge")
                    if merge:
                        row_span = merge.get("row_span", 1)
                        col_span = merge.get("col_span", 1)
                        if row_span > 1 or col_span > 1:
                            end_row = min(row_idx + row_span - 1, rows - 1)
                            end_col = min(col_idx + col_span - 1, cols - 1)
                            cell.merge(table.cell(end_row, end_col))

        # 표 테두리
        borders = table_def.get("borders")
        if borders:
            self._apply_table_borders(table, borders)

    def _apply_cell_shading(self, cell, color_hex):
        """셀 배경색 적용 (XML 직접 조작)"""
        if not color_hex:
            return
        color = color_hex.lstrip("#")
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), color)
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        cell._tc.get_or_add_tcPr().append(shading)

    def _apply_table_borders(self, table, borders_def):
        """표 테두리 설정 (XML 직접 조작)"""
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)

        borders = OxmlElement("w:tblBorders")
        edge_map = {
            "top": "top",
            "bottom": "bottom",
            "left": "left",
            "right": "right",
            "inside_h": "insideH",
            "inside_v": "insideV",
        }
        style_map = {
            "none": "none",
            "single": "single",
            "double": "double",
            "dotted": "dotted",
            "dashed": "dashed",
            "thick": "thick",
        }

        for json_key, xml_name in edge_map.items():
            edge_def = borders_def.get(json_key, {})
            if not edge_def or edge_def.get("style") == "none":
                continue
            el = OxmlElement(f"w:{xml_name}")
            el.set(
                qn("w:val"),
                style_map.get(edge_def.get("style", "single"), "single"),
            )
            el.set(qn("w:sz"), str(int(edge_def.get("width_pt", 1) * 8)))
            el.set(qn("w:space"), "0")
            color = edge_def.get("color", "#000000")
            el.set(qn("w:color"), color.lstrip("#") if color else "000000")
            borders.append(el)

        tblPr.append(borders)

    def _add_image(self, image_def):
        """이미지 삽입"""
        img_path = image_def.get("image_file", "")
        # 상대 경로를 base_dir 기준으로 해석
        if not os.path.isabs(img_path):
            img_path = os.path.join(self.base_dir, img_path)

        if not os.path.exists(img_path):
            print(f"경고: 이미지 파일 없음: {img_path}")
            return

        width = cm_val(image_def.get("width_cm"))
        height = cm_val(image_def.get("height_cm"))

        if width and height:
            self.doc.add_picture(img_path, width=width, height=height)
        elif width:
            self.doc.add_picture(img_path, width=width)
        else:
            self.doc.add_picture(img_path)

    def _setup_headers_footers(self):
        """머리글/바닥글 설정"""
        sections = self.template.get("sections", [])
        for section_def in sections:
            idx = section_def.get("index", 0)
            if idx >= len(self.doc.sections):
                continue
            section = self.doc.sections[idx]

            # 머리글
            headers_def = section_def.get("headers", {})
            if headers_def.get("primary"):
                header = section.header
                header.is_linked_to_previous = False
                paras = headers_def["primary"].get("paragraphs", [])
                for i, para_def in enumerate(paras):
                    if i == 0 and header.paragraphs:
                        p = header.paragraphs[0]
                    else:
                        p = header.add_paragraph()
                    p.text = para_def.get("text", "")
                    if para_def.get("style"):
                        try:
                            p.style = self.doc.styles[para_def["style"]]
                        except KeyError:
                            pass

            # 바닥글
            footers_def = section_def.get("footers", {})
            if footers_def.get("primary"):
                footer = section.footer
                footer.is_linked_to_previous = False
                paras = footers_def["primary"].get("paragraphs", [])
                for i, para_def in enumerate(paras):
                    if i == 0 and footer.paragraphs:
                        p = footer.paragraphs[0]
                    else:
                        p = footer.add_paragraph()
                    p.text = para_def.get("text", "")


# ─── 메인 실행 ──────────────────────────────────────────

def main():
    if len(sys.argv) < 3:
        print("사용법:")
        print("  Mode A: python generate_doc.py <template.json> <output.docx>")
        print("  Mode B: python generate_doc.py <template.json> <content.json> <output.docx>")
        sys.exit(1)

    generator = DocGenerator()

    if len(sys.argv) == 3:
        # Mode A: 원본 재현
        template_path = sys.argv[1]
        output_path = sys.argv[2]
        print(f"Mode A: 원본 재현")
        print(f"템플릿: {template_path}")
    else:
        # Mode B: 새 콘텐츠 적용
        template_path = sys.argv[1]
        content_path = sys.argv[2]
        output_path = sys.argv[3]
        print(f"Mode B: 새 콘텐츠 적용")
        print(f"템플릿: {template_path}")
        print(f"콘텐츠: {content_path}")
        generator.load_content(content_path)

    generator.load_template(template_path)
    print("템플릿 로드 완료")

    generator.generate()
    print("문서 생성 완료")

    saved_path = generator.save(output_path)
    print(f"저장 완료: {saved_path}")


if __name__ == "__main__":
    main()
